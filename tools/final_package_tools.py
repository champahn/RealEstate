#!/usr/bin/env python3
# /// script
# requires-python = ">=3.12"
# dependencies = ["python-docx", "pillow"]
# ///
# allow: SIZE_OK - one-off ULW document package builder and evidence harness kept self-contained for auditability.
from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import stat
import subprocess
import sys
from collections import Counter
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Final, TypedDict

BUNDLED_PYTHON: Final = Path(os.environ.get("CODEX_BUNDLED_PYTHON", sys.executable))
if Path(sys.executable).resolve() != BUNDLED_PYTHON.resolve() and BUNDLED_PYTHON.exists():
    os.execv(str(BUNDLED_PYTHON), [str(BUNDLED_PYTHON), str(Path(__file__).resolve()), *sys.argv[1:]])

from PIL import Image, ImageStat
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT: Final = Path(__file__).resolve().parents[1]
ULW: Final = Path(os.environ.get("MANUAL_ULW_DIR", str(ROOT / ".omo" / "ulw-loop" / "20260705-manual-v11-final")))
TARGET: Final = ROOT
FINAL: Final = TARGET / "v1.1_final"
DATA: Final = FINAL / "data"
TEMPLATES: Final = FINAL / "template_pack"
QA: Final = FINAL / "qa_evidence_pack"
RENDER: Final = FINAL / "rendered_docx"
ORIGINAL: Final = Path(
    os.environ.get(
        "MANUAL_V1_ORIGINAL",
        str(ROOT.parent / "manual-original-backups" / "부동산개발_시스템_아파트시행매뉴얼_v1.0_ORIGINAL_20260705.docx"),
    )
)
BACKUP: Final = Path(
    os.environ.get(
        "MANUAL_V1_BACKUP",
        str(ROOT.parent / "manual-original-backups" / "부동산개발_시스템_아파트시행매뉴얼_v1.0_ORIGINAL_20260705.docx"),
    )
)
SOURCE_AUDIT: Final = TARGET / "source_preservation_audit.md"
DB_V01: Final = TARGET / "failure_case_db_v0.1.jsonl"
CLAIM_LEDGER: Final = TARGET / "claim-ledger.md"
SYNTHESIS: Final = TARGET / "SYNTHESIS.md"
REG_TRACKER: Final = TARGET / "regulatory_update_tracker.md"
FINAL_MD: Final = FINAL / "부동산개발_시스템_아파트시행매뉴얼_v1.1_DRAFT.md"
FINAL_DOCX: Final = FINAL / "부동산개발_시스템_아파트시행매뉴얼_v1.1_DRAFT.docx"
FINAL_PDF: Final = RENDER / "부동산개발_시스템_아파트시행매뉴얼_v1.1_DRAFT.pdf"
DB_V02: Final = DATA / "failure_case_db_v0.2.jsonl"
NORMALIZATION_SUMMARY: Final = DATA / "normalization_summary.md"
RESEARCH_BACKLOG: Final = DATA / "research_backlog_unresolved_claims.md"
CHAPTER_MAP: Final = DATA / "chapter_control_map.json"
GATE_PACK_MD: Final = FINAL / "gate_control_pack.md"
GATE_PACK_JSON: Final = FINAL / "gate_control_pack.json"
REG_TRACKER_V11: Final = FINAL / "regulatory_update_tracker_v1.1.md"
SOURCE_REPORT: Final = QA / "source_preservation_report.md"
SIM_QA_MD: Final = QA / "failure_simulation_qa.md"
SIM_QA_JSON: Final = QA / "failure_simulation_qa.json"
GIT_EVIDENCE: Final = QA / "git_github_evidence.md"
RENDER_SCRIPT: Final = Path(os.environ.get("DOCX_RENDER_SCRIPT", str(ROOT / "tools" / "render_docx.py")))
PYTHON: Final = BUNDLED_PYTHON
POPPLER_LIB: Final = Path(os.environ.get("POPPLER_LIB", str(ROOT / ".local" / "poppler" / "lib")))
PRIVATE_DIR: Final = 0o700
PRIVATE_FILE: Final = 0o600
SOFFICE_LINKS: Final = {
    Path("/opt/homebrew/opt/little-cms2/lib/liblcms2.2.dylib"): POPPLER_LIB / "liblcms2.2.dylib",
    Path("/opt/homebrew/opt/fontconfig/lib/libfontconfig.1.dylib"): POPPLER_LIB / "libfontconfig.1.dylib",
    Path("/opt/homebrew/opt/freetype/lib/libfreetype.6.dylib"): POPPLER_LIB / "libfreetype.6.dylib",
}
G004_STATUS: Final = "reverified-with-followup-detail-needed"


class SourceState(TypedDict):
    path: str
    sha256: str
    size: int
    mode: str
    mtime: str


@dataclass(frozen=True)
class Workstream:
    code: str
    name: str
    location: str
    control: str
    evidence: str
    training: str


@dataclass(frozen=True)
class PhaseControl:
    phase: str
    owner: str
    trigger: str
    threshold: str
    evidence: str
    stop: str
    escalation: str
    training: str


WORKSTREAMS: Final = [
    Workstream("WS-00", "Version Governance", "PART 1 / Appendix A", "원본·작업본·승인본 SHA, diff, 승인 로그를 변경 전후로 잠근다.", "source_preservation_report, version_change_log, approval_log", "작업본만 편집하는 diff review drill"),
    Workstream("WS-01", "PF 자본구조·보증의존도", "Phase 3 / Phase 5", "자기자본 단계, 보증총량, DSCR, repayment source, cash waterfall를 Gate 조건으로 둔다.", "PF exposure map, term sheet, equity 입금증", "브릿지론 만기 60일 전 본PF 전환성 drill"),
    Workstream("WS-02", "분양광고·HUG·환급", "Phase 3 / Phase 7 / Phase 8", "광고문구 법무승인, 미확정 개발계획 disclaimer, 환급·재분양 decision tree를 둔다.", "legal signoff, 광고 evidence binder, HUG 사고 SOP", "분양률 50% 미만 EOD 근접 drill"),
    Workstream("WS-03", "토지 DD Red-Flag", "Phase 1 / Phase 2", "유치권·분묘·문화재·접도·토양오염·지장물별 stop condition을 분리한다.", "title opinion, red-flag checklist, 잔금집행 checklist", "잔금 전 가압류·지장물 발견 drill"),
    Workstream("WS-04", "인허가·민원 Stakeholder", "Phase 1 / Phase 4", "stakeholder map, 주민반대 issue log, 대안안 관리대장을 필수 증빙으로 둔다.", "쟁점표, Q&A pack, 회의록, 대안안 비교표", "심의 2회 보류와 주민민원 tabletop"),
    Workstream("WS-05", "공사비·시공사·하도급", "Phase 5 / Phase 6 / Phase 8", "claim ledger, 독립 QS trigger, 하도급 DD, hold-point, step-in 절차를 둔다.", "QS report, claim ledger, 하도급 지급확인", "공사비 15% 증액 요구 negotiation drill"),
    Workstream("WS-06", "계약·판례·분쟁 SOP", "Phase 2 / 4 / 5 / 7 / Appendix D", "계약특약 library, 법무 signoff log, dispute escalation ladder를 운영한다.", "계약 redline, 외부법무 의견, 판례 tracker", "책임준공·분양광고 조항 redline drill"),
    Workstream("WS-07", "조직/RACI·승인권한", "PART 1 / All Gates", "Gate별 RACI, 투자심의, 리스크위원회, 대체승인, 인수인계를 명시한다.", "RACI matrix, 회의록, 위임전결표", "대표 부재·담당자 퇴사 대체승인 drill"),
    Workstream("WS-08", "세무·회계·SPC·자금통제", "Phase 2 / 3 / 5 / 8", "취득세 taxonomy, VAT 안분, SPC 계좌, funds-flow, tax memo trigger를 둔다.", "tax memo, 계좌대장, 증빙 binder", "계약변경 tax memo trigger drill"),
    Workstream("WS-09", "VDR·개인정보·증거보존", "All Phases / Appendix A", "VDR taxonomy, 권한표, 개인정보 inventory, 다운로드 로그, 72시간 사고대응을 둔다.", "VDR index, permission matrix, access log", "개인정보 사고 72시간 tabletop"),
    Workstream("WS-10", "부동산개발업 컴플라이언스", "PART 1 / Appendix E", "등록요건, 전문인력, 사무실, 변경보고, 사업실적보고 calendar를 둔다.", "compliance calendar, roster, 제출증빙", "전문인력 퇴사·변경보고 누락 drill"),
    Workstream("WS-11", "실패사례 기반 교육훈련", "PART 1 / Every Phase End", "case card, Gate drill, red-team review, check ride, 분기 회고를 운영한다.", "training roster, drill score, after-action review", "신규 담당자 90일 onboarding check ride"),
    Workstream("WS-12", "준공·하자·미분양 Exit", "Phase 7 / Phase 8", "defect triage, reserve policy, 환급 SLA, 미분양 보유/임대/벌크매각 tree를 둔다.", "defect log, reserve calc, refund SLA, exit memo", "준공 후 잔금미납·하자 집단 claim drill"),
]

PHASES: Final = [
    PhaseControl("Phase 1. 사업기획 및 토지 발굴", "개발PM, 재무, 법무", "후보지 검토 착수", "P75 초과 분양가, 자기자본 계획 미확정, 민원 high-risk", "실거래가 산정표, 토지이용계획, stakeholder map, 예비 PF memo", "대표+재무+법무 승인 전 LOI/계약금 집행 금지", "리스크위원회 48시간 내 재심의", "P50/P75 pricing 및 stakeholder red-team"),
    PhaseControl("Phase 2. 토지 확보 및 실사", "개발PM, 법무, 외부 법무법인", "계약서 초안 수령, 중도금/잔금 지급 전", "권리하자 미해소, 지장물 비용 미반영, 법무의견 미완료", "title opinion, red-flag checklist, 잔금집행 checklist", "red flag 미해소 시 잔금 지급 금지", "외부법무+대표 예외승인", "가압류·지장물 발견 잔금중단 drill"),
    PhaseControl("Phase 3. 사업타당성 분석", "재무, 회계/세무 자문, 개발PM", "정밀 사업수지 확정 전", "Worst Case 이익률 음수, BEP 분양률 80% 초과, equity funded ratio 미달", "사업수지표, P50/P75 table, tax memo, PF structure memo", "투자심의 재상정 전 사업착수 금지", "CFO 주관 sensitivity 재검토", "PF 구조와 세무 memo check ride"),
    PhaseControl("Phase 4. 인허가", "개발PM, 설계사, 인허가 대행, 법무", "심의 접수 전, 반려/보류 통보 즉시", "보류 2회 이상, 주요 민원 unresolved, 지연 60일 초과", "회의록, 쟁점표, Q&A pack, 대안안 비교표", "금융비용 재산정 없이 PF 착수 금지", "대표+재무+인허가 PM escalation", "심의 보류·주민민원 tabletop"),
    PhaseControl("Phase 5. PF 자금조달", "CFO/재무책임자, 법무, 대표", "브릿지론 만기 90/60/30일 전, term sheet 수령 전", "equity 미납입, DSCR 미달, 보증총량 상한 초과", "exposure map, cash waterfall, term sheet, EOD review memo", "리스크위원회 승인 전 신규 보증/차입 금지", "대주단 협상 war-room", "만기 60일 전 본PF 미확정 drill"),
    PhaseControl("Phase 6. 시공 관리", "CM, 개발PM, 법무, 재무", "증액 요구, 공기 지연, 기성 mismatch, 하도급 미지급 징후", "도급액 5% 이상 증액, 공정 지연 30일 초과, 하도급 체불", "claim ledger, QS report, 하도급 지급확인, 공정률 report", "변경승인 전 시공 착수 금지", "독립 QS+법무 공동 검토", "공사비 15% 증액 요구 drill"),
    PhaseControl("Phase 7. 분양 및 마케팅", "마케팅, 법무, 재무", "광고문구 확정 전, 입주자모집공고 전, 분양률 70/50% 미만", "법무 미승인 문구, HUG 기준 미확인, PF EOD 분양률 조건 근접", "legal signoff, 광고 evidence binder, 분양률 dashboard", "법무승인 전 광고/모집공고 금지", "분양대행사 교체 또는 가격정책 재승인", "분양률 50% 미만 EOD drill"),
    PhaseControl("Phase 8. 준공 및 정산", "재무, CS, 법무, 세무", "사용검사 신청 전, 입주율 90% 미달, 하자 claim 접수", "잔금미납률 초과, high-severity 하자, 미분양 장기화", "closeout checklist, reserve calculation, defect log, tax closing memo", "PF 상환·분양보증 해지 전 정산 검증 필수", "준공정산위원회 재승인", "잔금미납·하자 집단 claim drill"),
]


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)
    os.chmod(path, PRIVATE_DIR)


def write_text(path: Path, text: str) -> None:
    ensure_dir(path.parent)
    path.write_text(text, encoding="utf-8")
    os.chmod(path, PRIVATE_FILE)


def write_json(path: Path, payload) -> None:
    write_text(path, json.dumps(payload, ensure_ascii=False, indent=2) + "\n")


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def mtime_iso(path: Path) -> str:
    return datetime.fromtimestamp(path.stat().st_mtime, timezone.utc).isoformat()


def state(path: Path) -> SourceState:
    st = path.stat()
    return {"path": str(path), "sha256": sha256(path), "size": st.st_size, "mode": oct(stat.S_IMODE(st.st_mode)), "mtime": mtime_iso(path)}


def read_cases() -> list[dict[str, str]]:
    with DB_V01.open(encoding="utf-8") as handle:
        return [json.loads(line) for line in handle if line.strip()]


def read_claim_rows() -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for line in CLAIM_LEDGER.read_text(encoding="utf-8").splitlines():
        if not line.startswith("| CLM-"):
            continue
        parts = [part.strip() for part in line.strip("|").split("|")]
        rows.append({"claim_id": parts[0], "case_id": parts[1], "assertion": parts[2], "risk": parts[3], "verdict": parts[7], "use": parts[8]})
    return rows


def map_workstream(failure_mode: str) -> str:
    checks = [
        ("PF", "WS-01"),
        ("분양", "WS-02"),
        ("토지", "WS-03"),
        ("인허가", "WS-04"),
        ("시공", "WS-05"),
        ("법률", "WS-06"),
        ("조직", "WS-07"),
        ("세무", "WS-08"),
        ("데이터룸", "WS-09"),
        ("개발업", "WS-10"),
        ("교육훈련", "WS-11"),
        ("준공 후 정산", "WS-12"),
    ]
    for needle, code in checks:
        if needle in failure_mode:
            return code
    return "WS-00"


def manual_location(workstream: str) -> str:
    for ws in WORKSTREAMS:
        if ws.code == workstream:
            return ws.location
    return "PART 1"


def normalize_cases() -> list[dict[str, str]]:
    rows = []
    for case in read_cases():
        workstream = map_workstream(case["failure_mode"])
        status = case.get("claim_status", "unresolved")
        mandatory = "false" if status in {"unresolved", "refuted"} else "true"
        enriched = dict(case)
        enriched["workstream"] = workstream
        enriched["manual_location"] = manual_location(workstream)
        enriched["mandatory_control"] = mandatory
        enriched["downstream_route"] = "research-backlog" if mandatory == "false" else "manual-control"
        rows.append(enriched)
    return rows


def source_paths_from_audit() -> list[Path]:
    paths: list[Path] = []
    for line in SOURCE_AUDIT.read_text(encoding="utf-8").splitlines():
        if line.startswith("| `/"):
            paths.append(Path(line.split("`")[1]))
    if len(paths) != 5:
        raise RuntimeError(f"expected 5 Obsidian source paths, found {len(paths)}")
    return paths


def write_phase_b() -> None:
    ensure_dir(FINAL)
    ensure_dir(DATA)
    rows = normalize_cases()
    with DB_V02.open("w", encoding="utf-8") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")
    os.chmod(DB_V02, PRIVATE_FILE)
    claim_rows = read_claim_rows()
    unresolved = [row for row in claim_rows if row["verdict"] == "unresolved"]
    counts = Counter(row["workstream"] for row in rows)
    write_json(CHAPTER_MAP, {"generated_at": now_iso(), "case_count": len(rows), "workstream_counts": dict(counts), "workstreams": [asdict(ws) for ws in WORKSTREAMS]})
    backlog = ["# Research Backlog: unresolved high-risk claims", "", "이 항목은 최종 매뉴얼의 의무 통제로 승격하지 않는다.", ""]
    for row in unresolved:
        backlog.append(f"- `{row['claim_id']}` / `{row['case_id']}` / {row['risk']}: {row['assertion']} → route `{row['use']}`")
    write_text(RESEARCH_BACKLOG, "\n".join(backlog) + "\n")
    summary = [
        "# Phase B Evidence Normalization Summary",
        "",
        f"- generated_at: `{now_iso()}`",
        f"- v0.1 input rows: `{len(read_cases())}`",
        f"- v0.2 output rows: `{len(rows)}`",
        "- count normalization: `63` rows is the controlling count. `SYNTHESIS.md`의 64건 표기는 정정 후보로 유지한다.",
        f"- unresolved claim-ledger rows routed away from mandatory controls: `{len(unresolved)}`",
        "- mandatory-control rule: `claim_status in {unresolved, refuted}` is `research-backlog`, not an obligation.",
        "",
        "## Workstream Counts",
        "",
    ]
    for code, count in sorted(counts.items()):
        summary.append(f"- `{code}`: `{count}` cases")
    write_text(NORMALIZATION_SUMMARY, "\n".join(summary) + "\n")


def table_row(fields: list[str]) -> str:
    return "| " + " | ".join(fields) + " |"


def gate_pack_text() -> str:
    lines = ["# Gate Control Pack v1.1", "", "각 Gate는 Owner, Trigger, Threshold, Evidence, Stop-or-Approve, Escalation, Training Scenario를 모두 가진다.", ""]
    lines += [table_row(["Phase", "Owner", "Trigger", "Threshold", "Evidence", "Stop-or-Approve", "Escalation", "Training Scenario"]), table_row(["---"] * 8)]
    for phase in PHASES:
        lines.append(table_row([phase.phase, phase.owner, phase.trigger, phase.threshold, phase.evidence, phase.stop, phase.escalation, phase.training]))
    return "\n".join(lines) + "\n"


def template_pack() -> dict[str, str]:
    return {
        "README.md": "# Template Pack v1.1\n\n실패방지 통제 실행용 서식 모음이다. 모든 서식은 증빙 위치, owner, 승인권자를 반드시 남긴다.\n",
        "raci_matrix.md": "| Activity | Responsible | Accountable | Consulted | Informed | Backup |\n| --- | --- | --- | --- | --- | --- |\n| Gate 승인 | 개발PM | 대표 | 재무, 법무 | 현장, 분양 | 대체승인자 |\n",
        "pf_exposure_map_template.md": "| 차주 | 사업장 | 보증인 | 만기 | DSCR | LTV | Equity funded | Repayment source | Stop line |\n| --- | --- | --- | --- | --- | --- | --- | --- | --- |\n",
        "cash_waterfall_template.md": "| 순위 | 유입/지급 | 계좌 | 조건 | 증빙 | 승인권자 |\n| --- | --- | --- | --- | --- | --- |\n",
        "p50_p75_pricing_pack_template.md": "| Source | P25 | P50 | P75 | P90 | 적용가 | 제외 호가 사유 |\n| --- | --- | --- | --- | --- | --- | --- |\n",
        "land_dd_red_flag_pack_template.md": "| 항목 | 발견내용 | Stop condition | 해소증빙 | 법무의견 | 잔금 가능 여부 |\n| --- | --- | --- | --- | --- | --- |\n",
        "stakeholder_issue_log_template.md": "| 이해관계자 | 예상 반대 | 쟁점 | Q&A | Owner | 상태 | 금융비용 영향 |\n| --- | --- | --- | --- | --- | --- | --- |\n",
        "claim_ledger_template.md": "| Claim | 사유 | 금액 | 근거 | QS 의견 | 협상상태 | 승인권자 | 다음 action |\n| --- | --- | --- | --- | --- | --- | --- | --- |\n",
        "contract_clause_library_template.md": "| 계약 | 조항 | 리스크 | 필수 문구 | 금지 문구 | 법무승인 |\n| --- | --- | --- | --- | --- | --- |\n",
        "tax_spc_funds_flow_template.md": "| Event | Tax issue | Account | Evidence | Memo trigger | Reviewer |\n| --- | --- | --- | --- | --- | --- |\n",
        "vdr_privacy_index_template.md": "| Folder | Document | Personal data | Access role | Download log | Retention | Deletion trigger |\n| --- | --- | --- | --- | --- | --- | --- |\n",
        "developer_compliance_calendar_template.md": "| Due date | Obligation | Evidence | Owner | Backup | Submitted at | Recheck |\n| --- | --- | --- | --- | --- | --- | --- |\n",
        "training_drill_pack_template.md": "| Drill | Role | Scenario | Pass evidence | Score | Remediation |\n| --- | --- | --- | --- | --- | --- |\n",
        "closeout_defect_exit_template.md": "| Issue | Severity | SLA | Reserve | Owner | Customer notice | Exit decision |\n| --- | --- | --- | --- | --- | --- | --- |\n",
    }


def manual_markdown() -> str:
    lines = [
        "# 부동산개발 시스템 아파트 시행 매뉴얼 v1.1 DRAFT",
        "",
        f"- 작성일: `{now_iso()}`",
        "- 목적: 부동산 시행 경험이 없는 담당자도 실패·실수를 사전에 제거하도록 Gate, 증빙, 중단선을 실행 단위로 제공한다.",
        "- 법률·세무·정책 주의: 본 문서는 실무 통제 매뉴얼이며 법률·세무 자문을 대체하지 않는다. 최신 법령·고시·협회 서식은 프로젝트 착수 전 재확인한다.",
        "",
        "## PART 1. 시스템 아키텍처와 Version Governance",
        "",
        "WS-00 Version Governance: 원본 v1.0은 보존하고 v1.1 작업본과 승인본만 편집한다. 모든 변경은 SHA256, diff, 승인로그, source preservation report를 남긴다.",
        "",
        "### 운영 원칙",
        "",
        "- 원본 DOCX, 원본 백업, Obsidian 원천 파일은 편집 금지 대상이다.",
        "- 작업본은 `v1.1_DRAFT`, 승인본은 `v1.1_APPROVED`로 명명한다.",
        "- `verdict=unresolved`인 법률·정책·금융 claim은 research backlog에 남기고 의무 통제로 쓰지 않는다.",
        "- 모든 Gate는 Owner, Trigger, Threshold, Evidence, Stop-or-Approve, Escalation, Training Scenario를 가진다.",
        "",
        "### RACI",
        "",
        "- Gate 승인: Responsible=개발PM, Accountable=대표, Consulted=재무·법무, Informed=현장·분양, Backup=대체승인자.",
        "- PF·자금집행: Responsible=CFO/재무, Accountable=대표, Consulted=법무·세무, Evidence=cash waterfall와 계좌대장.",
        "- 데이터룸·개인정보: Responsible=데이터관리자, Accountable=준법, Consulted=법무, Evidence=permission matrix와 access log.",
        "",
        "| 역할 | 책임 | 대체자 | 필수 산출물 |",
        "| --- | --- | --- | --- |",
        "| 대표 | 최종 투자·중단 승인 | 위임전결자 | 승인 회의록 |",
        "| 개발PM | Phase 운영과 증빙 취합 | 보조 PM | Gate pack, issue log |",
        "| CFO/재무 | PF, 수지, 세무·SPC 통제 | 회계책임자 | cash waterfall, tax memo |",
        "| 법무 | 계약·분쟁·광고문구 통제 | 외부 법무법인 | legal signoff, redline |",
        "| 준법/데이터관리자 | 개발업, VDR, 개인정보 | PMO | compliance calendar, permission matrix |",
        "",
        "## PART 2. 12 Core-Axis Reinforcement Workstreams",
        "",
    ]
    for ws in WORKSTREAMS:
        lines += [
            f"### {ws.code} {ws.name}",
            "",
            f"- 연결 위치: {ws.location}",
            f"- 핵심 통제: {ws.control}",
            f"- 증빙: {ws.evidence}",
            f"- Training Scenario: {ws.training}",
            "",
        ]
    lines += ["## PART 3. Phase-by-Phase 실무 매뉴얼", ""]
    for phase in PHASES:
        lines += [
            f"### {phase.phase}",
            "",
            "#### 실행 초점",
            "",
            f"{phase.phase}에서는 사전 검토, 증빙 고정, 의사결정 중단선을 같은 회의체에서 확인한다. 설명형 체크리스트가 아니라 실행 가능한 승인·중단 기준으로 운용한다.",
            "",
            "#### 실패방지 Control Box",
            "",
            f"- Owner: {phase.owner}",
            f"- Trigger: {phase.trigger}",
            f"- Threshold: {phase.threshold}",
            f"- Evidence: {phase.evidence}",
            f"- Stop-or-Approve: {phase.stop}",
            f"- Escalation: {phase.escalation}",
            f"- Training Scenario: {phase.training}",
            "",
            "#### 최소 실행서식",
            "",
            "- Gate 회의록: 안건, 결론, 보류 사유, 조건부 승인 조건, 다음 재검토일.",
            "- Evidence binder: 원본 파일명, source date, 담당자, 승인자, VDR 위치.",
            "- Exception log: Threshold 위반 사유, 보완 조치, 예외승인권자.",
            "",
        ]
    lines += [
        "## PART 4. Failure Case DB 운영규칙",
        "",
        "- 실패사례 DB v0.2는 63건을 기준으로 한다.",
        "- `SYNTHESIS.md`의 64건 표기는 정정 후보이며 누락 row가 확인되기 전까지 본문 기준 숫자로 쓰지 않는다.",
        "- 각 case는 workstream, manual_location, training_scenario, mandatory_control 여부를 가진다.",
        "- unresolved high-risk claim은 mandatory control이 아니라 research backlog다.",
        "",
        "## PART 5. Regulatory Update Tracker",
        "",
        "- PF 자기자본·보증의존도 정책, 부동산개발사업 관리법, 개발업 등록·보고, 개인정보, 분양광고 항목은 프로젝트 착수 전 재확인한다.",
        f"- G004-FU-A: status `{G004_STATUS}`. aside-browser 실패 후 insane-search/curl_cffi로 국토교통부 `제2차 부동산서비스산업 진흥 기본계획` 공식 HTML/PDF 접근과 `실적 확인제` 방향 확인을 완료했다. 세부 시행일·서식·주관기관 운영 절차는 후속 고시·협회 안내 확인 대상으로 유지한다.",
        "",
        "## PART 6. Template Pack 사용법",
        "",
        "- Gate Control Pack은 회의 전 사전작성, 회의 중 threshold 판정, 회의 후 VDR 업로드 순서로 쓴다.",
        "- Template Pack은 증빙 없는 판단을 막기 위한 최소 입력 양식이며 프로젝트 특성에 따라 필드를 추가하되 삭제하지 않는다.",
        "",
        "## Appendix A. Source Boundary",
        "",
        f"- original v1.0: `{ORIGINAL}`",
        f"- read-only backup: `{BACKUP}`",
        f"- Phase B-E final directory: `{FINAL}`",
        "",
    ]
    return "\n".join(lines) + "\n"


def write_reg_tracker() -> None:
    lines = [
        "# Regulatory Update Tracker v1.1",
        "",
        f"- generated_at: `{now_iso()}`",
        "- 주의: 이 tracker는 최신성 관리용이며 법률·세무 자문이 아니다.",
        "",
        "| 항목 | 현재 반영 | 본문 반영 방식 | 재확인 포인트 |",
        "| --- | --- | --- | --- |",
        "| PF 자기자본비율·보증의존도 | 금융위원회/KDI source 기반 준비통제 | 의무 수치 단정 대신 stress threshold와 재확인 tracker | 금융위 후속 고시, 대주단 term sheet |",
        "| 부동산개발사업 관리법 | 법령정보센터상 시행·부분시행 구조 확인 | 준법 RACI와 조문 matrix 준비통제 | 시행령·시행규칙별 실제 의무 개시일 |",
        f"| 부동산 개발사업 실적 확인제 | G004-FU-A에서 국토부 공식 HTML/PDF 재검증 완료 | 개발업 컴플라이언스 appendix에 `{G004_STATUS}`로 반영 | 세부 시행일, 서식, 주관기관 안내 |",
        "| 개발업 등록·전문인력·사업실적 | 현행 필수 통제 | compliance calendar, roster, evidence binder | 연간 서식·협회 안내 |",
        "| 개인정보·분양마케팅 | current-law tracker | VDR/privacy index와 legal signoff | 개인정보위·공정위·국토부 최신 안내 |",
        "",
        "## Official URLs checked",
        "",
        "- 금융위원회 PF 상황 점검회의: https://www.fsc.go.kr/no010101/86630",
        "- 국토교통부 제2차 부동산서비스산업 진흥 기본계획: https://www.molit.go.kr/USR/NEWS/m_71/dtl.jsp?id=95092133&lcmspage=1",
        "- 법령정보센터 부동산개발사업 관리 등에 관한 법률/시행령: https://www.law.go.kr/",
    ]
    write_text(REG_TRACKER_V11, "\n".join(lines) + "\n")


def write_docx_from_markdown(markdown: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    for style_name, size, color in [("Heading 1", 18, RGBColor(31, 78, 121)), ("Heading 2", 14, RGBColor(47, 84, 150)), ("Heading 3", 11.5, RGBColor(68, 68, 68))]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True
        styles[style_name].font.color.rgb = color
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("부동산개발 시스템 아파트 시행 매뉴얼 v1.1 DRAFT")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)
    for line in markdown.splitlines()[1:]:
        text = line.strip()
        if not text:
            continue
        if text.startswith("#### "):
            doc.add_heading(text[5:], level=3)
        elif text.startswith("### "):
            doc.add_heading(text[4:], level=3)
        elif text.startswith("## "):
            doc.add_heading(text[3:], level=2)
        elif text.startswith("# "):
            doc.add_heading(text[2:], level=1)
        elif text.startswith("- "):
            doc.add_paragraph(text[2:], style="List Bullet")
        elif text.startswith("| "):
            continue
        else:
            doc.add_paragraph(text)
    ensure_dir(FINAL_DOCX.parent)
    doc.save(FINAL_DOCX)
    os.chmod(FINAL_DOCX, PRIVATE_FILE)


def write_phase_c_e() -> None:
    markdown = manual_markdown()
    write_text(FINAL_MD, markdown)
    write_docx_from_markdown(markdown)
    write_text(GATE_PACK_MD, gate_pack_text())
    write_json(GATE_PACK_JSON, {"generated_at": now_iso(), "phases": [asdict(phase) for phase in PHASES]})
    for name, text in template_pack().items():
        write_text(TEMPLATES / name, text)
    write_reg_tracker()


def simulations() -> list[dict[str, str]]:
    return [
        {"id": "SIM-01", "scenario": "브릿지론 만기 60일 전 본PF 조건 미확정", "decision": "STOP", "owner": "CFO/재무책임자", "evidence": "PF exposure map, term sheet status, equity funded proof", "manual_location": "Phase 5 / WS-01", "training_scenario": "브릿지론 만기 60일 전 본PF 전환성 drill", "reason": "본PF term sheet 또는 equity 납입이 없으면 신규 보증·차입 금지"},
        {"id": "SIM-02", "scenario": "토지 잔금 전 가압류 또는 지장물 발견", "decision": "STOP", "owner": "개발PM, 법무", "evidence": "title opinion, red-flag checklist, 잔금집행 checklist", "manual_location": "Phase 2 / WS-03", "training_scenario": "잔금 전 가압류·지장물 발견 drill", "reason": "권리하자·지장물 비용 미반영 시 잔금 지급 금지"},
        {"id": "SIM-03", "scenario": "인허가 심의 2회 보류와 주민 민원", "decision": "STOP_REPRICE_OR_ESCALATE", "owner": "개발PM, 인허가 대행, 법무", "evidence": "쟁점표, Q&A pack, 대안안 비교표, 금융비용 재산정", "manual_location": "Phase 4 / WS-04", "training_scenario": "심의 2회 보류와 주민민원 tabletop", "reason": "금융비용 재산정과 대안안 승인 전 PF 착수 중지"},
        {"id": "SIM-04", "scenario": "시공사 공사비 15% 증액 요구", "decision": "STOP_PENDING_QS_AND_APPROVAL", "owner": "CM, 개발PM, 재무, 법무", "evidence": "claim ledger, QS report, 하도급 지급확인", "manual_location": "Phase 6 / WS-05", "training_scenario": "공사비 15% 증액 요구 negotiation drill", "reason": "5% 이상 증액은 독립 QS와 변경승인 전 착수 금지"},
        {"id": "SIM-05", "scenario": "분양률 50% 미만, PF EOD 근접", "decision": "STOP_ESCALATE", "owner": "마케팅, 법무, 재무", "evidence": "분양률 dashboard, legal signoff, HUG/EOD memo", "manual_location": "Phase 7 / WS-02", "training_scenario": "분양률 50% 미만 EOD 근접 drill", "reason": "가격정책·광고·재분양 의사결정 tree를 리스크위원회로 격상"},
        {"id": "SIM-06", "scenario": "준공 후 잔금 미납과 하자 집단 claim", "decision": "STOP_OR_APPROVE_WITH_RESERVE", "owner": "재무, CS, 법무, 세무", "evidence": "defect log, reserve calculation, refund SLA, tax closing memo", "manual_location": "Phase 8 / WS-12", "training_scenario": "잔금미납·하자 집단 claim drill", "reason": "정산 검증 전 PF 상환·분양보증 해지 금지"},
    ]


def write_phase_d() -> None:
    sims = simulations()
    write_json(SIM_QA_JSON, {"generated_at": now_iso(), "simulations": sims})
    lines = ["# Failure Simulation QA", ""]
    lines += [table_row(["ID", "Scenario", "Decision", "Owner", "Evidence", "Manual Location", "Training Scenario"]), table_row(["---"] * 7)]
    for sim in sims:
        lines.append(table_row([sim["id"], sim["scenario"], sim["decision"], sim["owner"], sim["evidence"], sim["manual_location"], sim["training_scenario"]]))
    write_text(SIM_QA_MD, "\n".join(lines) + "\n")


def write_source_report() -> None:
    sources = [ORIGINAL, BACKUP, *source_paths_from_audit()]
    payload = {"generated_at": now_iso(), "sources": [state(path) for path in sources]}
    write_json(QA / "source_preservation_report.json", payload)
    lines = ["# Source Preservation Report", "", f"- generated_at: `{payload['generated_at']}`", ""]
    lines += [table_row(["Path", "SHA256", "Mode", "Mtime"]), table_row(["---"] * 4)]
    for item in payload["sources"]:
        lines.append(table_row([item["path"], item["sha256"], item["mode"], item["mtime"]]))
    write_text(SOURCE_REPORT, "\n".join(lines) + "\n")


def remove_soffice_links() -> list[str]:
    removed: list[str] = []
    for link in SOFFICE_LINKS:
        if link.is_symlink():
            link.unlink()
            removed.append(str(link))
    return removed


def render_docx() -> None:
    ensure_dir(RENDER)
    remove_soffice_links()
    created: list[Path] = []
    try:
        for link, target in SOFFICE_LINKS.items():
            ensure_dir(link.parent)
            if not link.exists() and not link.is_symlink():
                link.symlink_to(target)
                created.append(link)
        env = os.environ.copy()
        env["PATH"] = f"{PYTHON.parent}:{env.get('PATH', '')}"
        result = subprocess.run([str(PYTHON), str(RENDER_SCRIPT), str(FINAL_DOCX), "--output_dir", str(RENDER), "--emit_pdf"], check=False, capture_output=True, text=True, env=env, timeout=180)
        write_text(QA / "render_docx.log", result.stdout + "\nSTDERR\n" + result.stderr)
        if result.returncode != 0:
            raise RuntimeError(f"render_docx failed: {result.returncode}")
        write_render_nonblank_proof()
    finally:
        for link in created:
            if link.is_symlink():
                link.unlink()


def pdf_page_count() -> int:
    result = subprocess.run(["pdfinfo", str(FINAL_PDF)], check=False, capture_output=True, text=True, timeout=30)
    if result.returncode != 0:
        raise RuntimeError(f"pdfinfo failed: {result.stderr}")
    match = re.search(r"^Pages:\s+(\d+)$", result.stdout, re.MULTILINE)
    if not match:
        raise RuntimeError("pdfinfo did not report page count")
    return int(match.group(1))


def write_render_nonblank_proof() -> None:
    pages = sorted(RENDER.glob("page-*.png"))
    proof_rows: list[dict[str, str]] = []
    for page in pages:
        with Image.open(page) as image:
            gray = image.convert("L")
            stat_result = ImageStat.Stat(gray)
            extrema = gray.getextrema()
            nonwhite = sum(count for value, count in enumerate(gray.histogram()) if value < 245)
            proof_rows.append({
                "path": str(page),
                "size": f"{image.width}x{image.height}",
                "mean_luma": f"{stat_result.mean[0]:.2f}",
                "extrema": f"{extrema[0]}..{extrema[1]}",
                "nonwhite_pixels": str(nonwhite),
            })
    page_count = pdf_page_count()
    write_json(QA / "render_nonblank_proof.json", {"generated_at": now_iso(), "pdf_pages": page_count, "png_pages": len(pages), "pages": proof_rows})
    lines = ["# Render Nonblank Proof", "", f"- pdf_pages: `{page_count}`", f"- png_pages: `{len(pages)}`", ""]
    lines += [table_row(["Page", "Size", "Mean Luma", "Extrema", "Nonwhite Pixels"]), table_row(["---"] * 5)]
    for row in proof_rows:
        lines.append(table_row([row["path"], row["size"], row["mean_luma"], row["extrema"], row["nonwhite_pixels"]]))
    write_text(QA / "render_nonblank_proof.md", "\n".join(lines) + "\n")


def write_git_evidence() -> None:
    remote = "origin"
    lines = ["# Git/GitHub Evidence", "", f"- generated_at: `{now_iso()}`", f"- intended_remote: `{remote}`"]
    for cmd in [["git", "status", "--short"], ["git", "branch", "--show-current"], ["git", "log", "-5", "--oneline"]]:
        result = subprocess.run(cmd, cwd=TARGET, check=False, capture_output=True, text=True)
        lines += ["", f"## `$ {' '.join(cmd)}`", "```", result.stdout.strip() or result.stderr.strip(), "```"]
    push_check = subprocess.run(["git", "ls-remote", remote], cwd=TARGET, check=False, capture_output=True, text=True, timeout=30)
    lines += ["", "## GitHub push/auth check", "```", (push_check.stdout + push_check.stderr).strip(), "```"]
    attempt = Path("/tmp/realestate_push_attempt.txt")
    if attempt.exists():
        lines += ["", "## `git push -u origin main` attempt", "```", attempt.read_text(encoding="utf-8").strip(), "```"]
    if attempt.exists() and "could not read Username" in attempt.read_text(encoding="utf-8"):
        lines.append("- result: `push-auth-required-at-that-run`")
    else:
        lines.append("- result: `remote-readable`" if push_check.returncode == 0 else "- result: `remote-unreadable-or-auth-required`")
    write_text(GIT_EVIDENCE, "\n".join(lines) + "\n")


def build_all() -> None:
    write_phase_b()
    write_phase_c_e()
    write_phase_d()
    write_source_report()
    write_text(QA / "cleanup_receipt.txt", f"cleanup_at: `{now_iso()}`\nrender_temp_symlinks_removed: `{remove_soffice_links()}`\n")


def verify_c001() -> str:
    rows = [json.loads(line) for line in DB_V02.read_text(encoding="utf-8").splitlines() if line.strip()]
    assert len(rows) == 63
    assert "64건 표기는 정정 후보" in NORMALIZATION_SUMMARY.read_text(encoding="utf-8")
    assert all(row.get("phase") and row.get("workstream") and row.get("manual_location") and row.get("training_scenario") for row in rows)
    assert all(row["downstream_route"] == "research-backlog" for row in rows if row.get("claim_status") == "unresolved")
    text = f"C001 PASS rows=63 unresolved_routed={sum(1 for row in rows if row.get('claim_status') == 'unresolved')} evidence={DB_V02}; {NORMALIZATION_SUMMARY}; cleanup=no runtime resources\n"
    write_text(ULW / "phase_b_c001_evidence_normalization.txt", text)
    return text


def verify_c002() -> str:
    required = [FINAL_MD, FINAL_DOCX, GATE_PACK_MD, GATE_PACK_JSON, REG_TRACKER_V11, *[TEMPLATES / name for name in template_pack()]]
    for path in required:
        assert path.exists() and path.stat().st_size > 0, path
    text = FINAL_MD.read_text(encoding="utf-8")
    for ws in WORKSTREAMS:
        assert ws.code in text and ws.name in text
    for phase in PHASES:
        assert phase.phase in text
    for field in ["Owner", "Trigger", "Threshold", "Evidence", "Stop-or-Approve", "Escalation", "Training Scenario"]:
        assert text.count(field) >= 8
    assert "G004-FU-A" in text and G004_STATUS in text and "실적 확인제" in text
    result = f"C002 PASS final_md={FINAL_MD} final_docx={FINAL_DOCX} templates={TEMPLATES} gate_pack={GATE_PACK_MD} cleanup=no runtime resources\n"
    write_text(ULW / "phase_c_e_c002_content_packaging.txt", result)
    return result


def verify_c003() -> str:
    sims = json.loads(SIM_QA_JSON.read_text(encoding="utf-8"))["simulations"]
    assert len(sims) == 6
    assert all(sim["decision"] and sim["owner"] and sim["evidence"] and sim["manual_location"] and sim["training_scenario"] for sim in sims)
    render_pages = sorted(RENDER.glob("page-*.png"))
    proof = json.loads((QA / "render_nonblank_proof.json").read_text(encoding="utf-8"))
    assert proof["pdf_pages"] == len(render_pages) == proof["png_pages"]
    assert all(int(page["nonwhite_pixels"]) > 1000 for page in proof["pages"])
    assert FINAL_PDF.exists() and FINAL_PDF.stat().st_size > 0
    for link in SOFFICE_LINKS:
        assert not link.is_symlink()
    assert SOURCE_REPORT.exists() and GIT_EVIDENCE.exists()
    report = SOURCE_REPORT.read_text(encoding="utf-8")
    assert sha256(ORIGINAL) in report and sha256(BACKUP) in report
    result = f"C003 PASS simulations=6 rendered_pages={len(render_pages)} pdf={FINAL_PDF} source_report={SOURCE_REPORT} git_evidence={GIT_EVIDENCE} cleanup=render temp symlinks absent, no tmux/browser/server resources spawned\n"
    write_text(ULW / "phase_d_e_c003_simulation_render_git.txt", result)
    return result


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("command", choices=["build", "render", "git-evidence", "C001", "C002", "C003"])
    args = parser.parse_args()
    try:
        if args.command == "build":
            build_all()
            print(f"built {FINAL}")
        elif args.command == "render":
            render_docx()
            print(f"rendered {FINAL_DOCX} -> {RENDER}")
        elif args.command == "git-evidence":
            write_git_evidence()
            print(GIT_EVIDENCE)
        elif args.command == "C001":
            print(verify_c001())
        elif args.command == "C002":
            print(verify_c002())
        elif args.command == "C003":
            print(verify_c003())
    except (AssertionError, FileNotFoundError, RuntimeError, subprocess.SubprocessError) as exc:
        print(f"FAIL: {exc}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

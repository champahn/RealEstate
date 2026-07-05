#!/usr/bin/env python3
# /// script
# requires-python = ">=3.12"
# dependencies = ["python-docx"]
# ///
# noqa: SIZE_OK - one-off audited DOCX repackage builder with large Korean control data constants.
from __future__ import annotations

import hashlib
import json
import os
import shutil
import xml.etree.ElementTree as ET
import zipfile
from dataclasses import asdict, dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Final

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT: Final = Path(__file__).resolve().parents[1]
FINAL: Final = ROOT / "v1.1_final"
BASE: Final = Path(
    os.environ.get(
        "READABLE_MANUAL_BASE",
        str(ROOT.parent / "manual-original-backups" / "부동산개발_시스템_아파트시행매뉴얼_v1.0_ORIGINAL_20260705.docx"),
    )
)
OLD_V11: Final = FINAL / "부동산개발_시스템_아파트시행매뉴얼_v1.1_DRAFT.docx"
OUT_DOCX: Final = FINAL / "부동산개발_시스템_아파트시행매뉴얼_v1.1_READABLE_DRAFT.docx"
OUT_MD: Final = FINAL / "부동산개발_시스템_아파트시행매뉴얼_v1.1_READABLE_DRAFT.md"
MANIFEST: Final = FINAL / "readable_rebuild_manifest.json"
BODY_FONT: Final = "AppleGothic"


@dataclass(frozen=True, slots=True)
class PhaseControl:
    phase: str
    owner: str
    trigger: str
    threshold: str
    evidence: str
    stop: str
    escalation: str
    training: str
    workstreams: str


@dataclass(frozen=True, slots=True)
class ChapterControl:
    chapter: str
    location: str
    insert_form: str
    required_terms: str


@dataclass(frozen=True, slots=True)
class WorkstreamAnchor:
    code: str
    name: str
    location: str
    control: str
    evidence: str
    training: str


PHASES: Final = [
    PhaseControl("Phase 1. 사업기획 및 토지 발굴", "개발PM, 재무, 법무", "후보지 검토 착수", "P75 초과 분양가, 자기자본 계획 미확정, 민원 high-risk", "실거래가 산정표, 토지이용계획, stakeholder map, 예비 PF memo", "대표+재무+법무 승인 전 LOI/계약금 집행 금지", "리스크위원회 48시간 내 재심의", "P50/P75 pricing 및 stakeholder red-team", "WS-03 토지 DD, WS-04 인허가·민원 stakeholder, WS-09 VDR"),
    PhaseControl("Phase 2. 토지 확보 및 실사 (Due Diligence)", "개발PM, 법무, 외부 법무법인", "계약서 초안 수령, 중도금/잔금 지급 전", "권리하자 미해소, 지장물 비용 미반영, 법무의견 미완료", "title opinion, red-flag checklist, 잔금집행 checklist, VDR index", "red flag 미해소 시 잔금 지급 금지", "외부법무+대표 예외승인", "가압류·지장물 발견 잔금중단 drill", "WS-03 토지 DD red-flag, WS-06 계약·판례 SOP, WS-08 세무·SPC, WS-09 VDR"),
    PhaseControl("Phase 3. 사업타당성 분석", "재무, 회계/세무 자문, 개발PM", "정밀 사업수지 확정 전", "Worst Case 이익률 음수, BEP 분양률 80% 초과, equity funded ratio 미달", "사업수지표, P50/P75 table, tax memo, PF structure memo", "투자심의 재상정 전 사업착수 금지", "CFO 주관 sensitivity 재검토", "PF 구조와 세무 memo check ride", "WS-01 PF 자본구조·보증의존도, WS-02 분양/HUG, WS-08 세무·회계·SPC"),
    PhaseControl("Phase 4. 인허가", "개발PM, 설계사, 인허가 대행, 법무", "심의 접수 전, 반려/보류 통보 즉시", "보류 2회 이상, 주요 민원 unresolved, 지연 60일 초과", "회의록, 쟁점표, Q&A pack, 대안안 비교표", "금융비용 재산정 없이 PF 착수 금지", "대표+재무+인허가 PM escalation", "심의 보류·주민민원 tabletop", "WS-04 인허가·민원 stakeholder, WS-06 계약·분쟁 SOP, WS-09 VDR"),
    PhaseControl("Phase 5. PF (Project Financing) 자금조달", "CFO/재무책임자, 법무, 대표", "브릿지론 만기 90/60/30일 전, term sheet 수령 전", "equity 미납입, DSCR 미달, 보증총량 상한 초과", "PF exposure map, cash waterfall, term sheet, EOD review memo", "리스크위원회 승인 전 신규 보증/차입 금지", "대주단 협상 war-room", "만기 60일 전 본PF 미확정 drill", "WS-01 PF 자본구조, WS-05 책임준공·신탁, WS-06 계약, WS-08 세무·SPC"),
    PhaseControl("Phase 6. 시공 관리", "CM, 개발PM, 법무, 재무", "증액 요구, 공기 지연, 기성 mismatch, 하도급 미지급 징후", "도급액 5% 이상 증액, 공정 지연 30일 초과, 하도급 체불", "claim ledger, QS report, 하도급 지급확인, 공정률 report", "변경승인 전 시공 착수 금지", "독립 QS+법무 공동 검토", "공사비 15% 증액 요구 negotiation drill", "WS-05 공사비·책임준공·신탁, WS-11 교육훈련"),
    PhaseControl("Phase 7. 분양 및 마케팅", "마케팅, 법무, 재무", "광고문구 확정 전, 입주자모집공고 전, 분양률 70/50% 미만", "법무 미승인 문구, HUG 기준 미확인, PF EOD 분양률 조건 근접", "legal signoff, 광고 evidence binder, 분양률 dashboard, HUG 사고 SOP, 개인정보 inventory", "법무승인 전 광고/모집공고 금지", "분양대행사 교체 또는 가격정책 재승인", "분양률 50% 미만 EOD drill 및 개인정보 사고 72시간 tabletop", "WS-02 분양광고·HUG·환급 SOP, WS-09 개인정보 사고대응, WS-12 Exit"),
    PhaseControl("Phase 8. 준공 및 정산", "재무, CS, 법무, 세무", "사용검사 신청 전, 입주율 90% 미달, 하자 claim 접수", "잔금미납률 초과, high-severity 하자, 미분양 장기화", "closeout checklist, reserve calculation, defect log, tax closing memo, refund SLA", "PF 상환·분양보증 해지 전 정산 검증 필수", "준공정산위원회 재승인", "잔금미납·하자 집단 claim drill", "WS-12 준공·하자·미분양 Exit, WS-08 세무 closing, WS-02 HUG 환급"),
]

CHAPTERS: Final = [
    ChapterControl("PF 자본구조와 보증의존도", "Phase 3 / Phase 5", "사업수지표와 PF Closing Gate에 삽입", "PF, 보증총량, DSCR, cash waterfall"),
    ChapterControl("토지 DD red-flag", "Phase 1 / Phase 2", "입지·실사 체크리스트에 stop condition 추가", "유치권, 분묘, 문화재, 접도, 토양오염, 지장물"),
    ChapterControl("인허가·민원 stakeholder", "Phase 1 / Phase 4", "Stakeholder map, issue log, Q&A pack 표로 추가", "주민반대, 쟁점표, 대안안"),
    ChapterControl("분양광고·HUG·환급 SOP", "Phase 3 / Phase 7 / Phase 8", "광고 법무승인과 환급 SLA Gate Evidence로 추가", "HUG, legal signoff, refund SLA"),
    ChapterControl("공사비·책임준공·신탁", "Phase 5 / Phase 6", "기존 책임준공형 관리형 토지신탁 장에 법률 risk rating 추가", "공사비, 책임준공, 신탁, QS"),
    ChapterControl("세무·회계·SPC", "Phase 2 / Phase 3 / Phase 5 / Phase 8", "tax memo trigger와 SPC 계좌대장 추가", "취득세, VAT, SPC, funds-flow"),
    ChapterControl("데이터룸·개인정보·증빙보존", "All Phases / Appendix E", "VDR index, permission matrix, access log, 72시간 사고대응 추가", "VDR, 개인정보, 개인정보 사고, 72시간, 사고대응"),
    ChapterControl("부동산개발업 컴플라이언스", "PART 1 / Appendix F", "등록·전문인력·사업실적보고 calendar 추가", "부동산개발업, 전문인력, 사업실적보고"),
    ChapterControl("준공 후 정산·하자·교육훈련", "Phase 8 / Appendix G", "defect triage와 training roster 추가", "하자, reserve, Training Scenario"),
]

WORKSTREAMS: Final = [
    WorkstreamAnchor("WS-00", "Version Governance", "PART 1 / Appendix A", "원본·작업본·승인본 SHA, diff, 승인 로그를 변경 전후로 잠근다.", "source_preservation_report, version_change_log, approval_log", "작업본만 편집하는 diff review drill"),
    WorkstreamAnchor("WS-01", "PF 자본구조·보증의존도", "Phase 3 / Phase 5", "자기자본 단계, 보증총량, DSCR, repayment source, cash waterfall를 Gate 조건으로 둔다.", "PF exposure map, term sheet, equity 입금증", "브릿지론 만기 60일 전 본PF 전환성 drill"),
    WorkstreamAnchor("WS-02", "분양광고·HUG·환급", "Phase 3 / Phase 7 / Phase 8", "광고문구 법무승인, 미확정 개발계획 disclaimer, 환급·재분양 decision tree를 둔다.", "legal signoff, 광고 evidence binder, HUG 사고 SOP", "분양률 50% 미만 EOD 근접 drill"),
    WorkstreamAnchor("WS-03", "토지 DD Red-Flag", "Phase 1 / Phase 2", "유치권·분묘·문화재·접도·토양오염·지장물별 stop condition을 분리한다.", "title opinion, red-flag checklist, 잔금집행 checklist", "잔금 전 가압류·지장물 발견 drill"),
    WorkstreamAnchor("WS-04", "인허가·민원 Stakeholder", "Phase 1 / Phase 4", "stakeholder map, 주민반대 issue log, 대안안 관리대장을 필수 증빙으로 둔다.", "쟁점표, Q&A pack, 회의록, 대안안 비교표", "심의 2회 보류와 주민민원 tabletop"),
    WorkstreamAnchor("WS-05", "공사비·시공사·하도급", "Phase 5 / Phase 6 / Phase 8", "claim ledger, 독립 QS trigger, 하도급 DD, hold-point, step-in 절차를 둔다.", "QS report, claim ledger, 하도급 지급확인", "공사비 15% 증액 요구 negotiation drill"),
    WorkstreamAnchor("WS-06", "계약·판례·분쟁 SOP", "Phase 2 / 4 / 5 / 7 / Appendix D", "계약특약 library, 법무 signoff log, dispute escalation ladder를 운영한다.", "계약 redline, 외부법무 의견, 판례 tracker", "책임준공·분양광고 조항 redline drill"),
    WorkstreamAnchor("WS-07", "조직/RACI·승인권한", "PART 1 / All Gates", "Gate별 RACI, 투자심의, 리스크위원회, 대체승인, 인수인계를 명시한다.", "RACI matrix, 회의록, 위임전결표", "대표 부재·담당자 퇴사 대체승인 drill"),
    WorkstreamAnchor("WS-08", "세무·회계·SPC·자금통제", "Phase 2 / 3 / 5 / 8", "취득세 taxonomy, VAT 안분, SPC 계좌, funds-flow, tax memo trigger를 둔다.", "tax memo, 계좌대장, 증빙 binder", "계약변경 tax memo trigger drill"),
    WorkstreamAnchor("WS-09", "VDR·개인정보·증거보존", "All Phases / Appendix E", "VDR taxonomy, 권한표, 개인정보 inventory, 다운로드 로그, 72시간 사고대응을 둔다.", "VDR index, permission matrix, access log", "개인정보 사고 72시간 tabletop"),
    WorkstreamAnchor("WS-10", "부동산개발업 컴플라이언스", "PART 1 / Appendix E", "등록요건, 전문인력, 사무실, 변경보고, 사업실적보고 calendar를 둔다.", "compliance calendar, roster, 제출증빙", "전문인력 퇴사·변경보고 누락 drill"),
    WorkstreamAnchor("WS-11", "실패사례 기반 교육훈련", "PART 1 / Every Phase End", "case card, Gate drill, red-team review, check ride, 분기 회고를 운영한다.", "training roster, drill score, after-action review", "신규 담당자 90일 onboarding check ride"),
    WorkstreamAnchor("WS-12", "준공·하자·미분양 Exit", "Phase 7 / Phase 8", "defect triage, reserve policy, 환급 SLA, 미분양 보유/임대/벌크매각 tree를 둔다.", "defect log, reserve calc, refund SLA, exit memo", "준공 후 잔금미납·하자 집단 claim drill"),
]

STATIC_TOC_LINES: Final = [
    ("PART 1. 시스템 아키텍처", "3", 0),
    ("1.1 3계층(Layer) 구조", "3", 1),
    ("1.2 2축 모듈 매트릭스", "3", 1),
    ("1.3 Gate 시스템", "3", 1),
    ("1.4~1.7 문서 활용·v1.1 운영 보강·WS Map", "4-5", 1),
    ("PART 2. 아파트 시행 실무 매뉴얼 v1.1 READABLE DRAFT", "7", 0),
    ("Phase 1. 사업기획 및 토지 발굴", "7", 1),
    ("Phase 2. 토지 확보 및 실사", "11", 1),
    ("Phase 3. 사업타당성 분석", "14", 1),
    ("Phase 4. 인허가", "18", 1),
    ("Phase 5. PF 자금조달", "21", 1),
    ("Phase 6. 시공 관리", "26", 1),
    ("Phase 7. 분양 및 마케팅", "30", 1),
    ("Phase 8. 준공 및 정산", "33", 1),
    ("PART 3. 부록", "36", 0),
    ("부록 A-D. 서류·용어·로드맵·법령", "36-39", 1),
    ("Appendix E-G. VDR·개인정보·컴플라이언스·교육훈련", "40-42", 1),
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def set_cell(cell_text: str, paragraph: Paragraph, *, bold: bool = False) -> None:
    paragraph.text = ""
    run = paragraph.add_run(cell_text)
    run.bold = bold
    run.font.size = Pt(8.5)
    run.font.name = BODY_FONT


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "B7C4D3")
        borders.append(tag)
    tbl_pr.append(borders)


def format_table(table: Table) -> None:
    table.autofit = True
    set_table_borders(table)
    for idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
            if idx == 0:
                shade_cell(cell, "E8EEF5")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def insert_paragraph_after(element, parent, text: str, style: str | None = None) -> Paragraph:
    node = OxmlElement("w:p")
    element.addnext(node)
    paragraph = Paragraph(node, parent)
    paragraph.style = style
    paragraph.text = text
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def insert_table_after(element, doc: DocxDocument, rows: list[list[str]]) -> Table:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for row_idx, values in enumerate(rows):
        for col_idx, value in enumerate(values):
            set_cell(value, table.cell(row_idx, col_idx).paragraphs[0], bold=row_idx == 0)
    element.addnext(table._tbl)
    format_table(table)
    return table


def paragraph_by_text(doc: DocxDocument, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise RuntimeError(f"paragraph not found: {text}")


def replace_cover_text(doc: DocxDocument) -> None:
    replacements = {
        "아파트 시행 실무 매뉴얼 v1.0": "아파트 시행 실무 매뉴얼 v1.1 READABLE DRAFT",
        "3-Layer Architecture  ×  8 Phase  ×  6 Gate System": "3-Layer Architecture × 8 Phase × Gate Evidence × Control Box",
        "PART 2. 아파트 시행 실무 매뉴얼 v1.0": "PART 2. 아파트 시행 실무 매뉴얼 v1.1 READABLE DRAFT",
    }
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)
        text = paragraph.text.strip()
        if text.startswith("v1.0 → 신탁사 책임준공"):
            paragraph.insert_paragraph_before("v1.1 READABLE → v1.0의 8 Phase·표·체크리스트 구조를 보존하고 v1.1 Control Box, Gate Evidence, Stop-or-Approve, RACI, Training Scenario를 각 Phase 안에 삽입")
    for section in doc.sections:
        containers = [
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ]
        for container in containers:
            for paragraph in container.paragraphs:
                if "아파트 시행 실무 매뉴얼 v1.0" in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "아파트 시행 실무 매뉴얼 v1.0",
                        "아파트 시행 실무 매뉴얼 v1.1 READABLE DRAFT",
                    )


def insert_part1_overlay(doc: DocxDocument) -> None:
    anchor = paragraph_by_text(doc, "PART 2. 아파트 시행 실무 매뉴얼 v1.1 READABLE DRAFT")
    anchor.paragraph_format.page_break_before = True
    parent = anchor._parent

    heading_15 = Paragraph(OxmlElement("w:p"), parent)
    heading_15.text = "1.5 v1.1 운영 보강 원칙"
    heading_15.style = "Heading 2"
    principle = Paragraph(OxmlElement("w:p"), parent)
    principle.text = "v1.1 READABLE DRAFT는 원본 v1.0의 8 Phase, 표, Gate 체크리스트를 보존하고, 신규 리스크 통제는 각 Phase 안에 표 형태로 삽입한다. 모든 Gate는 Owner/RACI, Trigger, Threshold, Gate Evidence, Stop-or-Approve, Escalation, Training Scenario를 가진다."
    raci_table = doc.add_table(rows=6, cols=3)
    for row_idx, values in enumerate([["Role", "RACI 책임", "Gate Evidence"], ["대표", "Accountable - 최종 투자·중단 승인", "승인 회의록, 예외승인 기록"], ["개발PM", "Responsible - Phase 운영과 증빙 취합", "Gate pack, issue log, VDR 위치"], ["CFO/재무", "Responsible - PF·자금·세무 통제", "cash waterfall, tax memo, 계좌대장"], ["법무", "Consulted - 계약·광고·분쟁 통제", "legal signoff, redline, 판례 tracker"], ["준법/데이터관리자", "Responsible - VDR·개인정보·개발업", "permission matrix, access log, compliance calendar"]]):
        for col_idx, value in enumerate(values):
            set_cell(value, raci_table.cell(row_idx, col_idx).paragraphs[0], bold=row_idx == 0)
    format_table(raci_table)

    heading_16 = Paragraph(OxmlElement("w:p"), parent)
    heading_16.text = "1.6 v1.1 Core Control Map"
    heading_16.style = "Heading 2"
    chapter_rows = [["Next Manual Chapter", "삽입 위치", "삽입 형식", "필수 키워드"], *[[c.chapter, c.location, c.insert_form, c.required_terms] for c in CHAPTERS]]
    chapter_table = doc.add_table(rows=len(chapter_rows), cols=len(chapter_rows[0]))
    for row_idx, values in enumerate(chapter_rows):
        for col_idx, value in enumerate(values):
            set_cell(value, chapter_table.cell(row_idx, col_idx).paragraphs[0], bold=row_idx == 0)
    format_table(chapter_table)

    heading_ws = Paragraph(OxmlElement("w:p"), parent)
    heading_ws.text = "1.7 v1.1 Workstream Anchor Map"
    heading_ws.style = "Heading 2"
    ws_rows = [["Code", "Workstream", "Manual Location", "Training Scenario"], *[[ws.code, ws.name, ws.location, ws.training] for ws in WORKSTREAMS]]
    ws_table = doc.add_table(rows=len(ws_rows), cols=len(ws_rows[0]))
    for row_idx, values in enumerate(ws_rows):
        for col_idx, value in enumerate(values):
            set_cell(value, ws_table.cell(row_idx, col_idx).paragraphs[0], bold=row_idx == 0)
    format_table(ws_table)

    nodes = [
        heading_15._p,
        principle._p,
        raci_table._tbl,
        heading_16._p,
        chapter_table._tbl,
        heading_ws._p,
        ws_table._tbl,
    ]
    cursor = anchor._p
    for node in reversed(nodes):
        cursor.addprevious(node)
        cursor = node


def insert_phase_controls(doc: DocxDocument) -> None:
    for phase in PHASES:
        anchor = paragraph_by_text(doc, phase.phase)
        parent = anchor._parent
        p1 = insert_paragraph_after(anchor._p, parent, "v1.1 실패방지 Control Box", "Heading 3")
        t1 = insert_table_after(p1._p, doc, [["구분", "내용"], ["Owner/RACI", phase.owner], ["Trigger", phase.trigger], ["Threshold", phase.threshold], ["Gate Evidence", phase.evidence], ["Stop-or-Approve", phase.stop], ["Escalation", phase.escalation], ["Training Scenario", phase.training], ["Related Workstreams", phase.workstreams]])
        p2 = insert_paragraph_after(t1._tbl, parent, "Gate Evidence & Stop-or-Approve 체크리스트", "Heading 3")
        insert_table_after(p2._p, doc, [["확인 항목", "필수 증빙", "판정"], ["증빙 원본성", "VDR 원본 파일명, source date, 승인자", "누락 시 STOP"], ["자금·수지 영향", "cash waterfall, tax memo, sensitivity", "Threshold 초과 시 재심의"], ["법무·계약 검토", "legal signoff, redline, 판례 tracker", "미승인 문구·조항 사용 금지"], ["교육훈련", phase.training, "분기 drill 또는 신규 담당자 check ride"]])


def append_new_appendices(doc: DocxDocument) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix E. v1.1 Evidence Binder & VDR 운영표", level=1)
    doc.add_paragraph("모든 Phase의 Gate Evidence는 VDR index, permission matrix, access log, source date를 가진다. 개인정보가 포함된 파일은 최소권한 원칙으로 접근권한을 분리하고 다운로드 로그를 보존한다. 개인정보 사고 또는 VDR 권한오류가 발견되면 72시간 내 사고대응 tabletop을 열고, 영향 범위·접근 로그·통지 필요성을 확인한다.")
    table = doc.add_table(rows=1 + len(PHASES), cols=4)
    for idx, value in enumerate(["Phase", "Evidence Binder", "VDR/개인정보", "보존 기준"]):
        set_cell(value, table.cell(0, idx).paragraphs[0], bold=True)
    for row_idx, phase in enumerate(PHASES, start=1):
        for col_idx, value in enumerate([phase.phase, phase.evidence, "permission matrix / access log", "Gate 승인본 기준 5년 이상 또는 법정기간"]):
            set_cell(value, table.cell(row_idx, col_idx).paragraphs[0])
    format_table(table)
    doc.add_heading("Appendix E-1. WS Anchor & 72시간 개인정보 사고대응", level=2)
    ws_table = doc.add_table(rows=1 + len(WORKSTREAMS), cols=5)
    for idx, value in enumerate(["Code", "Workstream", "Location", "Control", "Training"]):
        set_cell(value, ws_table.cell(0, idx).paragraphs[0], bold=True)
    for row_idx, ws in enumerate(WORKSTREAMS, start=1):
        for col_idx, value in enumerate([ws.code, ws.name, ws.location, ws.control, ws.training]):
            set_cell(value, ws_table.cell(row_idx, col_idx).paragraphs[0])
    format_table(ws_table)
    incident_table = doc.add_table(rows=5, cols=4)
    incident_rows = [
        ["Trigger", "Owner", "72시간 사고대응 Evidence", "Stop-or-Approve"],
        ["개인정보 사고, VDR 오권한, 외부 공유 링크 노출", "준법/데이터관리자, 법무, 개발PM", "access log, 다운로드 로그, 개인정보 inventory, 영향범위 memo", "72시간 내 사고대응 회의 전 외부 공유 중단"],
        ["대주단·시공사·분양대행사 자료 반출", "개발PM, 법무", "permission matrix, 반출승인서, 회수확인", "승인 없는 재공유 STOP"],
        ["개인정보 사고 72시간 tabletop", "준법/데이터관리자", "tabletop 시나리오, 대응시각표, 통지 판단 기록", "훈련 미실시 시 다음 Gate 보류"],
        ["사고대응 종료", "대표, 법무", "사후보고서, 재발방지 task owner", "대표 승인 후 정상화"],
    ]
    for row_idx, values in enumerate(incident_rows):
        for col_idx, value in enumerate(values):
            set_cell(value, incident_table.cell(row_idx, col_idx).paragraphs[0], bold=row_idx == 0)
    format_table(incident_table)
    doc.add_heading("Appendix F. 부동산개발업 컴플라이언스 Calendar", level=1)
    doc.add_paragraph("등록요건, 전문인력 roster, 사무실 요건, 변경보고, 사업실적보고는 준법 담당자가 calendar로 관리한다. 전문인력 퇴사·사무실 이전·사업실적보고 누락은 즉시 대표와 개발PM에게 escalation한다.")
    doc.add_heading("Appendix G. 교육훈련과 준공 후 정산 운영", level=1)
    doc.add_paragraph("실패사례 기반 교육훈련은 각 Phase 종료 회고와 분기 drill로 운영한다. 준공 후에는 defect triage, reserve calculation, refund SLA, tax closing memo를 확인하기 전 PF 상환·분양보증 해지를 승인하지 않는다.")


def write_markdown() -> None:
    lines = ["# 부동산개발 시스템 아파트 시행 매뉴얼 v1.1 READABLE DRAFT", "", "v1.0 원본의 8 Phase 구조와 표/체크리스트를 보존하고 v1.1 통제 보강 요소를 각 Phase 안에 삽입한 개선본이다.", "", "## Phase Control Overlay", ""]
    for phase in PHASES:
        lines += [f"### {phase.phase}", "", "| Field | Value |", "| --- | --- |", f"| Owner/RACI | {phase.owner} |", f"| Trigger | {phase.trigger} |", f"| Threshold | {phase.threshold} |", f"| Gate Evidence | {phase.evidence} |", f"| Stop-or-Approve | {phase.stop} |", f"| Escalation | {phase.escalation} |", f"| Training Scenario | {phase.training} |", f"| Related Workstreams | {phase.workstreams} |", ""]
    lines += ["## Next Manual Chapter Coverage", "", "| Chapter | Location | Insert Form | Required Terms |", "| --- | --- | --- | --- |"]
    lines += [f"| {item.chapter} | {item.location} | {item.insert_form} | {item.required_terms} |" for item in CHAPTERS]
    lines += ["", "## Workstream Anchor Map", "", "| Code | Workstream | Location | Control | Evidence | Training |", "| --- | --- | --- | --- | --- | --- |"]
    lines += [f"| {ws.code} | {ws.name} | {ws.location} | {ws.control} | {ws.evidence} | {ws.training} |" for ws in WORKSTREAMS]
    OUT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")
    os.chmod(OUT_MD, 0o600)


def docx_text_and_tables(path: Path) -> tuple[str, int]:
    w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    with zipfile.ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    text = "\n".join("".join(t.text or "" for t in p.iter(w + "t")) for p in root.iter(w + "p"))
    return text, len(list(root.iter(w + "tbl")))


def patch_docx_xml_text(path: Path, replacements: dict[str, str]) -> None:
    temp_path = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                xml_text = data.decode("utf-8")
                for old, new in replacements.items():
                    xml_text = xml_text.replace(old, new)
                data = xml_text.encode("utf-8")
            target.writestr(item, data)
    temp_path.replace(path)
    os.chmod(path, 0o600)


def patch_docx_fonts(path: Path) -> None:
    temp_path = path.with_suffix(".tmp.docx")
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    font_attrs = ("ascii", "hAnsi", "eastAsia", "cs")
    ET.register_namespace("w", w_ns)
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                root = ET.fromstring(data)
                changed = False
                for rfonts in root.iter(f"{{{w_ns}}}rFonts"):
                    for attr in font_attrs:
                        key = f"{{{w_ns}}}{attr}"
                        if rfonts.get(key) != BODY_FONT:
                            rfonts.set(key, BODY_FONT)
                            changed = True
                    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
                        key = f"{{{w_ns}}}{attr}"
                        if key in rfonts.attrib:
                            del rfonts.attrib[key]
                            changed = True
                if changed:
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            target.writestr(item, data)
    temp_path.replace(path)
    os.chmod(path, 0o600)


def scrub_docx_metadata(path: Path) -> None:
    temp_path = path.with_suffix(".tmp.docx")
    ns = {
        "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
        "dc": "http://purl.org/dc/elements/1.1/",
        "dcterms": "http://purl.org/dc/terms/",
    }
    ET.register_namespace("cp", ns["cp"])
    ET.register_namespace("dc", ns["dc"])
    ET.register_namespace("dcterms", ns["dcterms"])
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "docProps/core.xml":
                root = ET.fromstring(data)
                creator = root.find("dc:creator", ns)
                if creator is not None:
                    creator.text = "RealEstate Manual Team"
                last_modified_by = root.find("cp:lastModifiedBy", ns)
                if last_modified_by is not None:
                    last_modified_by.text = "RealEstate Manual Team"
                revision = root.find("cp:revision", ns)
                if revision is not None:
                    revision.text = "1"
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            target.writestr(item, data)
    temp_path.replace(path)
    os.chmod(path, 0o600)


def make_xml_paragraph(text: str) -> ET.Element:
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    paragraph = ET.Element(f"{{{w_ns}}}p")
    run = ET.SubElement(paragraph, f"{{{w_ns}}}r")
    text_node = ET.SubElement(run, f"{{{w_ns}}}t")
    text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_node.text = text
    return paragraph


def make_xml_page_break() -> ET.Element:
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    paragraph = ET.Element(f"{{{w_ns}}}p")
    run = ET.SubElement(paragraph, f"{{{w_ns}}}r")
    br = ET.SubElement(run, f"{{{w_ns}}}br")
    br.set(f"{{{w_ns}}}type", "page")
    return paragraph


def format_static_toc_line(title: str, page: str, level: int) -> str:
    prefix = "    " * level
    dot_count = max(8, 72 - len(prefix) - len(title) - len(page))
    return f"{prefix}{title}{'.' * dot_count}{page}"


def replace_docx_toc_static(path: Path) -> None:
    temp_path = path.with_suffix(".tmp.docx")
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ET.register_namespace("w", w_ns)
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename == "word/document.xml":
                root = ET.fromstring(data)
                body = root.find(f"{{{w_ns}}}body")
                if body is None:
                    raise RuntimeError("document body not found")
                children = list(body)

                def child_text(child: ET.Element) -> str:
                    return "".join(node.text or "" for node in child.iter(f"{{{w_ns}}}t"))

                start = next(
                    idx
                    for idx, child in enumerate(children)
                    if child_text(child).startswith("PART 1. 시스템 아키텍처")
                    and child_text(child) != "PART 1. 시스템 아키텍처"
                )
                end = next(
                    idx
                    for idx, child in enumerate(children[start + 1 :], start=start + 1)
                    if child_text(child) == "PART 1. 시스템 아키텍처"
                )
                for child in children[start:end]:
                    body.remove(child)

                static_nodes = [
                    make_xml_paragraph(format_static_toc_line(title, page, level))
                    for title, page, level in STATIC_TOC_LINES
                ]
                static_nodes.append(make_xml_page_break())
                for offset, node in enumerate(static_nodes):
                    body.insert(start + offset, node)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            target.writestr(item, data)
    temp_path.replace(path)
    os.chmod(path, 0o600)


def write_manifest() -> None:
    text, tables = docx_text_and_tables(OUT_DOCX)

    def display_path(path: Path) -> str:
        try:
            return str(path.relative_to(ROOT))
        except ValueError:
            try:
                return str(path.relative_to(ROOT.parent))
            except ValueError:
                return path.name

    payload = {
        "generated_at": datetime.now(UTC).isoformat(),
        "base_docx": display_path(BASE),
        "base_sha256": sha256(BASE),
        "old_v11_docx": display_path(OLD_V11),
        "old_v11_sha256": sha256(OLD_V11),
        "out_docx": display_path(OUT_DOCX),
        "out_sha256": sha256(OUT_DOCX),
        "out_md": display_path(OUT_MD),
        "out_docx_tables": tables,
        "out_docx_chars": len(text),
        "phase_count": sum(1 for idx in range(1, 9) if f"Phase {idx}" in text),
        "required_term_counts": {term: text.count(term) for term in ["PF", "HUG", "VDR", "RACI", "Gate Evidence", "Stop-or-Approve", "Training Scenario", "책임준공", "신탁"]},
        "phase_controls": [asdict(phase) for phase in PHASES],
        "chapter_controls": [asdict(chapter) for chapter in CHAPTERS],
        "workstream_anchors": [asdict(workstream) for workstream in WORKSTREAMS],
    }
    MANIFEST.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    os.chmod(MANIFEST, 0o600)


def main() -> int:
    if not BASE.exists():
        raise FileNotFoundError(BASE)
    FINAL.mkdir(parents=True, exist_ok=True)
    shutil.copy2(BASE, OUT_DOCX)
    os.chmod(OUT_DOCX, 0o600)
    doc = Document(str(OUT_DOCX))
    replace_cover_text(doc)
    insert_part1_overlay(doc)
    insert_phase_controls(doc)
    append_new_appendices(doc)
    doc.save(OUT_DOCX)
    patch_docx_xml_text(
        OUT_DOCX,
        {"아파트 시행 실무 매뉴얼 v1.0": "아파트 시행 실무 매뉴얼 v1.1 READABLE DRAFT"},
    )
    replace_docx_toc_static(OUT_DOCX)
    patch_docx_fonts(OUT_DOCX)
    scrub_docx_metadata(OUT_DOCX)
    write_markdown()
    write_manifest()
    print(f"wrote {OUT_DOCX}")
    print(f"wrote {OUT_MD}")
    print(f"wrote {MANIFEST}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
